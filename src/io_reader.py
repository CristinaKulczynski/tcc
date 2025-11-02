from __future__ import annotations
from pathlib import Path
from typing import Dict, Any
from docx import Document
import fitz  # PyMuPDF

SUPPORTED_EXTENSIONS = {".docx", ".pdf"}


def validate_supported_file_type(file_path: Path) -> None:
    if file_path.suffix.lower() not in SUPPORTED_EXTENSIONS:
        raise ValueError(f"Formato não suportado: {file_path.suffix}. Use .docx ou .pdf")
    if not file_path.is_file():
        raise FileNotFoundError(str(file_path))


# -------- DOCX --------
def extract_docx_text(file_path: Path) -> str:
    doc = Document(str(file_path))
    parts = []
    for p in doc.paragraphs:
        if p.text:
            parts.append(p.text)
    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                if c.text.strip():
                    parts.append(c.text.strip())
    for s in doc.sections:
        for h in s.header.paragraphs:
            if h.text:
                parts.append(h.text)
        for f in s.footer.paragraphs:
            if f.text:
                parts.append(f.text)
    return "\n".join(parts)


def count_docx_images(file_path: Path) -> int:
    doc = Document(str(file_path))
    rel = getattr(doc.part, "related_parts", {})
    return sum(1 for x in rel.values() if getattr(x, "content_type", "").startswith("image/"))


# -------- PDF --------
def extract_pdf_text(file_path: Path) -> str:
    with fitz.open(str(file_path)) as pdf:
        return "\n".join(page.get_text("text") for page in pdf)


def count_pdf_images(file_path: Path) -> int:
    with fitz.open(str(file_path)) as pdf:
        xrefs = set()
        for page in pdf:
            for img in page.get_images(full=True):
                xrefs.add(img[0])
        return len(xrefs)


def get_pdf_pages(file_path: Path) -> int:
    with fitz.open(str(file_path)) as pdf:
        return pdf.page_count


# -------- Orquestrador --------
def extract_document_properties(file_path: Path) -> Dict[str, Any]:
    validate_supported_file_type(file_path)
    ext = file_path.suffix.lower()

    if ext == ".docx":
        text = extract_docx_text(file_path)
        images = count_docx_images(file_path)
        pages = None
    else:
        text = extract_pdf_text(file_path)
        images = count_pdf_images(file_path)
        pages = get_pdf_pages(file_path)

    # Contagem simples: percorre o texto caractere por caractere
    total_chars = 0
    total_chars_no_space = 0
    for ch in text:
        total_chars += 1
        if ch not in (" ", "\u00A0"):  # ignora apenas espaços normais e NBSP
            total_chars_no_space += 1

    return {
        "characters_count": total_chars,
        "characters_no_space_count": total_chars_no_space,
        "pages_count": pages,
        "images_count": images,
    }
