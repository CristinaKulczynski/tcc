from pathlib import Path
from typing import Dict, Any
import docx
import xml.etree.ElementTree as ET

def validar_arquivo_docx(caminho_arquivo: Path) -> None:
    if caminho_arquivo.suffix.lower() != ".docx" or not caminho_arquivo.is_file():
        raise ValueError(f"Arquivo inválido: {caminho_arquivo}")

def ler_texto_docx(caminho_arquivo: Path) -> str:
    documento = docx.Document(str(caminho_arquivo))
    partes_texto = []

    for paragrafo in documento.paragraphs:
        if paragrafo.text.strip():
            partes_texto.append(paragrafo.text.strip())

    for tabela in documento.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                if celula.text.strip():
                    partes_texto.append(celula.text.strip())

    return "\n".join(partes_texto)

def contar_imagens_docx(caminho_arquivo: Path) -> int:
    documento = docx.Document(str(caminho_arquivo))
    count = 0
    # procura imagens em parágrafos
    for paragrafo in documento.paragraphs:
        count += len(paragrafo._element.xpath('.//w:drawing | .//w:pict'))
    # procura imagens dentro de tabelas
    for tabela in documento.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                count += len(celula._element.xpath('.//w:drawing | .//w:pict'))
    return count

def obter_paginas_appxml(app_xml_path: Path) -> int | None:
    if not app_xml_path.exists():
        return None
    try:
        tree = ET.parse(app_xml_path)
        root = tree.getroot()
        for elem in root.iter():
            if elem.tag.endswith("Pages"):
                return int(elem.text)
    except Exception:
        return None
    return None

def extract_document_properties_docx(caminho_arquivo: Path) -> Dict[str, Any]:
    validar_arquivo_docx(caminho_arquivo)

    texto = ler_texto_docx(caminho_arquivo)
    imagens = contar_imagens_docx(caminho_arquivo)

    app_xml_path = Path(__file__).resolve().parent.parent / "app.xml"
    paginas = obter_paginas_appxml(app_xml_path)

    texto_sem_quebras = texto.replace("\n", "").replace("\r", "")
    caracteres_total = len(texto_sem_quebras)
    caracteres_sem_espaco = sum(1 for caractere in texto_sem_quebras if not caractere.isspace())

    print("Texto extraído do documento:\n")
    print(texto)

    return {
        "characters_count": caracteres_total,
        "characters_no_space_count": caracteres_sem_espaco,
        "pages_count": paginas,
        "images_count": imagens,
        "text": texto
    }
